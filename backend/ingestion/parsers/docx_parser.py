"""DOCX Parser — extracts text from Word documents."""

import io
from docx import Document
from ingestion import TextChunk
from utils.logger import logger


def parse_docx(content: bytes, filename: str) -> list[TextChunk]:
    """Parse a DOCX file and return text chunks per section."""
    chunks = []
    try:
        doc = Document(io.BytesIO(content))
        logger.info(f"Parsing DOCX: {filename}")

        current_text = []
        current_section = "Document"
        page_estimate = 1

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Detect headings as section breaks
            if para.style and para.style.name.startswith("Heading"):
                # Save previous section
                if current_text:
                    chunks.append(TextChunk(
                        text="\n".join(current_text),
                        page_number=page_estimate,
                        section=current_section,
                        metadata={"source": filename},
                    ))
                    current_text = []
                current_section = text
                page_estimate += 1
            else:
                current_text.append(text)

        # Save last section
        if current_text:
            chunks.append(TextChunk(
                text="\n".join(current_text),
                page_number=page_estimate,
                section=current_section,
                metadata={"source": filename},
            ))

        # Also extract table content
        for table in doc.tables:
            rows_text = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    rows_text.append(" | ".join(cells))
            if rows_text:
                chunks.append(TextChunk(
                    text="\n".join(rows_text),
                    page_number=0,
                    section="Table",
                    metadata={"source": filename, "type": "table"},
                ))

        logger.info(f"DOCX parsed: {len(chunks)} sections")
    except Exception as e:
        logger.error(f"DOCX parsing failed for {filename}: {e}")
        raise

    return chunks
